# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
from random import randint
from random import choice
from docx import Document
from docx.shared import Inches
import png




s = ['0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0111110111110111110111110111110111110111110111110111110111110',
     '0000000000000000000000000000000000000000000000000000000000000',]



document = Document()
fieldside = 0
fieldside = int(input())
numero = 0
numero = int(input())

def letter(s: str):
    if s == "d1":
        return "N"
    elif s == "d2":
        return  "E"
    elif s == "d3":
        return "S"
    else:
        return "W"

numbero = 0
while numbero < numero:
     numbero +=1
     s = []
     for x in range(5):
          s.append((('0' + '00000')*fieldside + '0')*5)
     for x in range(fieldside):
          for x in range(25):
               s.append(('0'*5 + '11111'*5)*fieldside + '0'*5)
          for x in range(5):
               s.append((('0' + '00000') * fieldside + '0')*5)


     s = [[int(c) for c in row] for row in s]


     Allvalues = [[y for y in range(2)] for x in range(fieldside**2)]
     valindex = 0
     for yy in range(fieldside):
          for xx in range(fieldside):
               Allvalues[valindex][0] = yy
               Allvalues[valindex][1] = xx
               valindex += 1


     minpx = 0
     minpy = 0
     maxpx = fieldside-1
     maxpy = fieldside-1


     dlen1 = 0
     dlen2 = 0
     dlen3 = 0
     setlen = 0
     setlen = fieldside
     dlen1 = randint(1,setlen-2)
     setlen = setlen - dlen1
     dlen2 = randint(1,setlen-1)
     setlen = setlen - dlen2
     dlen3 = setlen
     d1 = randint(1,4)
     d2 = randint(1,4)
     while d1 == d2 or abs(d2-d1) == 2:
          d2 = randint(1,4)
     d3 = randint(1,4)
     while d3 == d2 or abs(d2-d3) == 2:
          d3 = randint(1,4)


     currloc = []
     answer = []
     checkpoint1 = []
     checkpoint2 = []
     checkpoint3 = []
     currloc.append(randint(0, fieldside-1))
     currloc.append(randint(0, fieldside-1))
     answer = currloc.copy()
     failsave = False
     savec = 0

     d1decided = False
     while d1decided == False:
          savec += 1
          d1 = randint(1, 4)
          if d1 == 1:
               if (currloc[0] - dlen1) > -1:
                    currloc[0] = currloc[0] - dlen1
                    checkpoint1 = currloc.copy()
                    d1decided = True

          if d1 == 2:
               if (currloc[1] + dlen1) < fieldside:
                    currloc[1] = currloc[1] + dlen1
                    checkpoint1 = currloc.copy()
                    d1decided = True
          if d1 == 3:
               if (currloc[0] + dlen1) < fieldside:
                    currloc[0] = currloc[0] + dlen1
                    checkpoint1 = currloc.copy()
                    d1decided = True
          if d1 == 4:
               if (currloc[1] - dlen1) > -1:
                    currloc[1] = currloc[1] - dlen1
                    checkpoint1 = currloc.copy()
                    d1decided = True
          if savec > 20:
               failsave = True
               break
     if failsave == True:
          print('BROKEN')
          numbero -= 1
          continue
     savec = 0
     d2decided = False
     while d2decided == False:
          savec += 1
          d2 = randint(1, 4)
          while d1 == d2 or abs(d2 - d1) == 2:
               d2 = randint(1, 4)
          if d2 == 1:
               if (currloc[0] - dlen2) > -1:
                    currloc[0] = currloc[0] - dlen2
                    checkpoint2 = currloc.copy()
                    d2decided = True

          if d2 == 2:
               if (currloc[1] + dlen2) < fieldside:
                    currloc[1] = currloc[1] + dlen2
                    checkpoint2 = currloc.copy()
                    d2decided = True
          if d2 == 3:
               if (currloc[0] + dlen2) < fieldside:
                    currloc[0] = currloc[0] + dlen2
                    checkpoint2 = currloc.copy()
                    d2decided = True
          if d2 == 4:
               if (currloc[1] - dlen2) > -1:
                    currloc[1] = currloc[1] - dlen2
                    checkpoint2 = currloc.copy()
                    d2decided = True
          if savec > 20:
               failsave = True
               break
     if failsave == True:
          print('BROKEN')
          numbero -= 1
          continue

     savec = 0
     d3decided = False
     while d3decided == False:
          savec += 1
          d3 = randint(1, 4)
          while d3 == d2 or abs(d2 - d3) == 2:
               d3 = randint(1, 4)
          if d3 == 1:
               if (currloc[0] - dlen3) > -1:
                    currloc[0] = currloc[0] - dlen3
                    checkpoint3 = currloc.copy()
                    d3decided = True

          if d3 == 2:
               if (currloc[1] + dlen3) < fieldside:
                    currloc[1] = currloc[1] + dlen3
                    checkpoint3 = currloc.copy()
                    d3decided = True
          if d3 == 3:
               if (currloc[0] + dlen3) < fieldside:
                    currloc[0] = currloc[0] + dlen3
                    checkpoint3 = currloc.copy()
                    d3decided = True
          if d3 == 4:
               if (currloc[1] - dlen3) > -1:
                    currloc[1] = currloc[1] - dlen3
                    checkpoint3 = currloc.copy()
                    d3decided = True
          if savec > 20:
               failsave = True
               break
     if failsave == True:
          print('BROKEN')
          numbero -= 1
          continue


     print(str(dlen1) + ' d' + str(d1))
     print(str(dlen2) + ' d' + str(d2))
     print(str(dlen3) + ' d' + str(d3))

     dirs = []
     dirs.append(str(dlen1) + ' d' + str(d1))
     dirs.append(str(dlen2) + ' d' + str(d2))
     dirs.append(str(dlen3) + ' d' + str(d3))

     print(answer)
     print(checkpoint1)
     print(checkpoint2)
     print(checkpoint3)

     NotRoute = Allvalues.copy()
     Route = []

     for x in [d1, d2, d3]:
          for a in [[answer, checkpoint1], [checkpoint1, checkpoint2], [checkpoint2, checkpoint3]]:
               if x == 1:
                    for x in range(a[0][0], a[1][0] , -1):
                         if [x, a[0][1]] in NotRoute:
                              NotRoute.remove([x, a[0][1]])
               if x == 2:
                    for x in range(a[0][1], a[1][1]):
                         if [a[0][0], x] in NotRoute:
                              NotRoute.remove([a[0][0], x])
               if x == 3:
                    for x in range(a[0][0],a[1][0]):
                         if [x, a[0][1]] in NotRoute:
                              NotRoute.remove([x, a[0][1]])
               if x == 4:
                    for x in range(a[0][1], a[1][1], -1):
                         if [a[0][0], x] in NotRoute:
                              NotRoute.remove([a[0][0], x])

     if checkpoint1 in NotRoute:
          NotRoute.remove(checkpoint1)
     if checkpoint2 in NotRoute:
          NotRoute.remove(checkpoint2)
     if checkpoint3 in NotRoute:
          NotRoute.remove(checkpoint3)

     for x in range(0, len(Allvalues)):
          if Allvalues[x] not in NotRoute:
               Route.append(Allvalues[x])



     lefttofill = Allvalues.copy()
     lefttofill.remove(answer)
     myrand = []
     mycurr = []
     blacks = []


     mycurr = choice(lefttofill)


     while lefttofill != []:
          weclear = True
          indexate = []
          myrand = []
          toremove = []
          cx = choice(lefttofill)
          indexate = cx.copy()
          myrand.append(indexate)
          for it in [[d1, dlen1], [d2, dlen2], [d3, dlen3]]:
               if it[0] == 1:
                    for x in range(1, it[1] + 1):
                         myrand.append([indexate[0] - x, indexate[1]])
                    indexate = [indexate[0] - x, indexate[1]]
               if it[0] == 2:
                    for x in range(1, it[1] + 1):
                         myrand.append([indexate[0], indexate[1] + x])
                    indexate = [indexate[0], indexate[1] + x]
               if it[0] == 3:
                    for x in range(1, it[1] + 1):
                         myrand.append([indexate[0] + x, indexate[1]])
                    indexate = [indexate[0] + x, indexate[1]]
               if it[0] == 4:
                    for x in range(1, it[1] + 1):
                         myrand.append([indexate[0], indexate[1] - x])
                    indexate = [indexate[0], indexate[1] - x]
          for x in myrand:
               if x[0] > fieldside - 1 or x[0] < 0:
                    weclear = False
               if x[1] > fieldside - 1 or x[1] < 0:
                    weclear = False
               if x in blacks:
                    weclear = False
               if x in Route:
                    toremove.append(x)
          for e in toremove:
               myrand.remove(e)
          if weclear == True:
               blacks.append(choice(myrand))

          lefttofill.remove(cx)



     ax = blacks[0][0]
     bx = blacks[0][1]
     an = 0
     bn = 0
     blocks = 0
     blocks = len(blacks)
     for x in range(0, len(blacks)):
          ax = blacks[x][0]
          bx = blacks[x][1]
          an = 0
          bn = 0
          while an < 25:
               while bn < 25:
                    (s[(5 + ax * 30 + an)][5 + bx * 30 + bn]) = 0
                    bn += 1
               bn = 0
               an += 1




     w = png.Writer(len(s[0]), len(s), greyscale=True, bitdepth=1)
     f = open('field.png', 'wb')
     w.write(f, s)
     f.close()






     document.add_picture('field.png', width=Inches(2.5))

     p = document.add_paragraph("Directions: ")
     s = ""


     s += letter(dirs[0].split()[1]) + dirs[0].split()[0] + " "
     s += letter(dirs[1].split()[1]) + dirs[1].split()[0] + " "
     s += letter(dirs[2].split()[1]) + dirs[2].split()[0] + " "

     p.add_run(s)





document.save('demo.docx')
